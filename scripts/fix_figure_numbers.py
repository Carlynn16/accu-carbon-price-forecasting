"""
Renumber figures in Statistical_Report.docx (in-place edit, no rebuild).

Mapping (old -> new):
  Figure 1  -> 1   Figure 2  -> 2
  Figure 6  -> 3   Figure 5  -> 4   Figure 7  -> 5
  Figure 3  -> 6   Figure 4  -> 7
  Figure 8  -> 8   Figure 9  -> 9
  Figure D1 -> 10  Figure D2 -> 11  Figure E1 -> 12  Figure E2 -> 13

Two-pass approach:
  Pass 1: old label -> unique @@PLACEHOLDER@@
  Pass 2: @@PLACEHOLDER@@ -> new label
Handles text split across paragraph runs without collapsing formatting.
"""

import re
import shutil
import sys
from pathlib import Path

from docx import Document

REPO   = Path(__file__).parent.parent
SRC    = REPO / "Statistical_Report.docx"
BACKUP = REPO / "Statistical_Report_backup.docx"
PDF    = REPO / "Statistical_Report.pdf"

# ── Two-pass substitution maps ────────────────────────────────────────────────
# Pass 1: existing label -> placeholder
# Listed most-specific first so longer patterns match before shorter ones.
PASS1 = {
    "Figure D1": "@@FIG10@@",
    "Figure D2": "@@FIG11@@",
    "Figure E1": "@@FIG12@@",
    "Figure E2": "@@FIG13@@",
    "Figure 3":  "@@FIG6@@",
    "Figure 4":  "@@FIG7@@",
    "Figure 5":  "@@FIG4@@",
    "Figure 6":  "@@FIG3@@",
    "Figure 7":  "@@FIG5@@",
}
# Pass 2: placeholder -> final label
PASS2 = {
    "@@FIG3@@":  "Figure 3",
    "@@FIG4@@":  "Figure 4",
    "@@FIG5@@":  "Figure 5",
    "@@FIG6@@":  "Figure 6",
    "@@FIG7@@":  "Figure 7",
    "@@FIG10@@": "Figure 10",
    "@@FIG11@@": "Figure 11",
    "@@FIG12@@": "Figure 12",
    "@@FIG13@@": "Figure 13",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _replace_in_text(text: str, old: str, new: str) -> str:
    """
    Replace old with new in text.
    For pure-numeric figures like 'Figure 3', add a lookahead so
    'Figure 3' does not match inside 'Figure 30' (if it existed).
    """
    if re.match(r"^Figure \d+$", old):
        return re.sub(re.escape(old) + r"(?!\d)", new, text)
    return text.replace(old, new)


def _present_in_text(text: str, old: str) -> bool:
    if re.match(r"^Figure \d+$", old):
        return bool(re.search(re.escape(old) + r"(?!\d)", text))
    return old in text


def _find_match_span(text: str, old: str):
    """Return (start, end) of the first occurrence of old in text, or None."""
    if re.match(r"^Figure \d+$", old):
        m = re.search(re.escape(old) + r"(?!\d)", text)
        return (m.start(), m.end()) if m else None
    idx = text.find(old)
    return (idx, idx + len(old)) if idx != -1 else None


def _apply_subs_to_para(para, subs: dict) -> bool:
    """
    Apply substitutions to a paragraph preserving per-run formatting.

    Strategy for each (old -> new):
      1. Per-run replacement (handles the common case of a single run).
      2. Cross-run fallback when the reference spans multiple runs:
         - Find the match position in the concatenated paragraph text.
         - Identify which runs overlap the match.
         - Rewrite the first affected run (prefix + new text),
           clear intermediate runs, rewrite the last run (suffix).
    """
    changed = False

    for old, new in subs.items():
        # ── 1. Per-run pass ───────────────────────────────────────────────
        for run in para.runs:
            replaced = _replace_in_text(run.text, old, new)
            if replaced != run.text:
                run.text = replaced
                changed = True

        # ── 2. Cross-run fallback ─────────────────────────────────────────
        # Repeat until all occurrences are resolved (rare: >1 cross-run hit)
        while True:
            full = _para_text(para)
            if not _present_in_text(full, old):
                break

            span = _find_match_span(full, old)
            if span is None:
                break
            match_start, match_end = span

            # Build run position map: [(run_text_start, run_text_end), ...]
            pos = 0
            run_spans = []
            for run in para.runs:
                run_spans.append((pos, pos + len(run.text)))
                pos += len(run.text)

            # Runs that overlap the match
            affected = [
                i for i, (rs, re_end) in enumerate(run_spans)
                if re_end > match_start and rs < match_end
            ]

            if len(affected) < 2:
                # Should not happen after per-run pass, but bail safely
                break

            fi, li = affected[0], affected[-1]
            first_rs = run_spans[fi][0]
            last_re  = run_spans[li][1]

            prefix = full[first_rs:match_start]   # text in first run before match
            suffix = full[match_end:last_re]       # text in last run after match

            runs = para.runs
            runs[fi].text = prefix + new
            for i in affected[1:-1]:               # clear intermediate runs
                runs[i].text = ""
            runs[li].text = suffix

            changed = True

    return changed


def _all_paragraphs(doc: Document):
    """Yield every paragraph in the document body and all table cells."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # 1. Backup
    shutil.copy2(SRC, BACKUP)
    print(f"Backup: {BACKUP.name}")

    # 2. Load
    doc = Document(SRC)

    # 3. Two-pass replacement
    pass1_hits = 0
    for para in _all_paragraphs(doc):
        if _apply_subs_to_para(para, PASS1):
            pass1_hits += 1

    pass2_hits = 0
    for para in _all_paragraphs(doc):
        if _apply_subs_to_para(para, PASS2):
            pass2_hits += 1

    print(f"Pass 1 paragraphs changed: {pass1_hits}")
    print(f"Pass 2 paragraphs changed: {pass2_hits}")

    # 4. Save
    doc.save(SRC)
    print(f"Saved:  {SRC.name}")

    # 5. Verify
    doc2 = Document(SRC)
    fig_re = re.compile(r"Figure\s+(?:\d+|[A-Z]\d+)")

    print("\n=== Figure references in document order ===")
    all_refs = []
    for para in _all_paragraphs(doc2):
        text = _para_text(para)
        for m in fig_re.finditer(text):
            all_refs.append(m.group())
            print(f"  {m.group()!r:20s}  | {text[:60]!r}")

    # Check captions specifically (look for "Figure N." or "Figure N " at start)
    cap_nums = []
    for para in _all_paragraphs(doc2):
        text = _para_text(para).strip()
        m = re.match(r"Figure\s+(\d+)", text)
        if m and (len(text) < 5 or text[m.end():m.end()+1] in (".", " ", "")):
            cap_nums.append(int(m.group(1)))

    cap_sorted = sorted(cap_nums)
    print(f"\nCaption figure numbers (in document order): {cap_nums}")
    print(f"Sorted:                                     {cap_sorted}")
    if cap_nums == cap_sorted and len(set(cap_nums)) == len(cap_nums):
        print("OK — captions are already sequential (reading order = numeric order).")
    else:
        # Duplicates / gaps
        from collections import Counter
        dupes = [n for n, c in Counter(cap_nums).items() if c > 1]
        expected = list(range(1, len(cap_nums) + 1))
        missing = sorted(set(expected) - set(cap_nums))
        if dupes:
            print(f"WARNING: duplicate caption numbers: {dupes}")
        if missing:
            print(f"WARNING: missing caption numbers: {missing}")

    # Check for any remaining old-style labels
    old_labels = ["Figure D1", "Figure D2", "Figure E1", "Figure E2",
                  "Figure 3", "Figure 4", "Figure 5", "Figure 6", "Figure 7"]
    remaining = []
    for para in _all_paragraphs(doc2):
        text = _para_text(para)
        for lbl in old_labels:
            if _present_in_text(text, lbl):
                remaining.append((lbl, text[:80]))
    if remaining:
        print("\nWARNING: unreplaced old labels:")
        for lbl, ctx in remaining:
            print(f"  {lbl!r} in {ctx!r}")
    else:
        print("OK — no unreplaced old labels found.")

    # 6. PDF conversion
    try:
        from docx2pdf import convert
        convert(str(SRC), str(PDF))
        print(f"\nPDF saved: {PDF.name}")
    except Exception as exc:
        print(f"\nPDF conversion failed ({type(exc).__name__}: {exc})")
        sys.exit(1)


if __name__ == "__main__":
    main()

"""
Build Statistical_Report.docx (and optionally Statistical_Report.pdf).

Re-running this script regenerates the document cleanly from scratch —
no duplicated sections.

Usage:
    python scripts/build_report.py
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt

from src.baselines import drift_predict, random_walk_predict
from src.eda import (
    compute_stationarity_tests,
    compute_staleness_stats,
    plot_acf_pacf,
    plot_actual_vs_pred,
    plot_error_by_regime,
    plot_feature_target_corr,
    plot_momentum,
    plot_price_timeline,
    plot_returns,
    plot_skill_by_horizon,
    plot_staleness,
    plot_target_distributions,
    plot_volatility,
)
from src.evaluate import build_results_table, compute_metrics, prepare_horizon
from src.features import build_features, save_features
from src.dl_models import run_dl_models
from src.explain import run_explain
from src.models import run_all_models
from src.significance import directional_accuracy_move_days, run_dm_tests
from src.report import (
    build_baselines_section,
    build_conclusion_section,
    build_data_section,
    build_dl_section,
    build_eda_section,
    build_explainability_section,
    build_features_section,
    build_intro,
    build_limitations_section,
    build_modeling_section,
    build_significance_section,
)

# ── Paths ─────────────────────────────────────────────────────────────────────
REPO       = Path(__file__).parent.parent
PROCESSED  = REPO / "data" / "processed"
FIGURES    = REPO / "figures"
DOCX_OUT   = REPO / "Statistical_Report.docx"
PDF_OUT    = REPO / "Statistical_Report.pdf"
TARGET_COL = "ACCU spot price (Generic)"


def _ensure_figures(
    train: pd.DataFrame,
    val:   pd.DataFrame,
    test:  pd.DataFrame,
    feat_train: pd.DataFrame,
    stale_stats: dict,
) -> None:
    """Generate all EDA figures (overwrites existing ones)."""
    FIGURES.mkdir(parents=True, exist_ok=True)
    plot_price_timeline(train, val, test, FIGURES / "fig_price_timeline.png")
    plot_target_distributions(train, val, test, FIGURES / "fig_target_dist_by_split.png")
    plot_returns(train, val, test, FIGURES / "fig_returns.png")
    plot_volatility(train, val, test, FIGURES / "fig_volatility.png")
    plot_acf_pacf(train, FIGURES / "fig_acf_pacf.png")
    plot_staleness(train, val, test, FIGURES / "fig_staleness.png", stale_stats=stale_stats)
    plot_momentum(train, FIGURES / "fig_momentum.png", stale_stats=stale_stats)
    plot_feature_target_corr(
        feat_train, target_col="target_1", outpath=FIGURES / "fig_feature_target_corr.png"
    )


def _compute_baselines(
    feat_val:   pd.DataFrame,
    feat_test:  pd.DataFrame,
    full_price: pd.Series,
    horizons:   tuple[int, ...] = (1, 7, 30),
) -> pd.DataFrame:
    """Run random-walk and drift baselines; return tidy results DataFrame."""
    records = []
    for h in horizons:
        for split_name, feat_df in [("val", feat_val), ("test", feat_test)]:
            data = prepare_horizon(feat_df, h)

            rw_pred    = random_walk_predict(data["n_rows"])
            rw_metrics = compute_metrics(rw_pred, data["y"], data["price_anchor"])
            rw_rmse    = rw_metrics["RMSE"]
            rw_mae     = rw_metrics["MAE"]
            rw_metrics.update({"RMSE_skill_%": 0.0, "MAE_skill_%": 0.0})
            records.append({"model": "random_walk", "horizon": h, "split": split_name,
                             **rw_metrics})

            d_pred    = drift_predict(data["dates"], full_price, h)
            d_metrics = compute_metrics(d_pred, data["y"], data["price_anchor"],
                                        rw_rmse=rw_rmse, rw_mae=rw_mae)
            records.append({"model": "drift", "horizon": h, "split": split_name,
                             **d_metrics})
    return build_results_table(records)


def _compute_baseline_preds(
    feat_val:   pd.DataFrame,
    feat_test:  pd.DataFrame,
    full_price: pd.Series,
    horizons:   tuple[int, ...] = (1, 7, 30),
) -> pd.DataFrame:
    """Return long-form predictions for RW and drift baselines."""
    records: list[dict] = []
    for h in horizons:
        for split_name, feat_df in [("val", feat_val), ("test", feat_test)]:
            data        = prepare_horizon(feat_df, h)
            rw_preds    = np.zeros(data["n_rows"])
            drift_preds = drift_predict(data["dates"], full_price, h)

            for d, anch, act, rw_p, dr_p in zip(
                data["dates"], data["price_anchor"], data["y"], rw_preds, drift_preds
            ):
                records.append({
                    "model": "random_walk", "horizon": h, "split": split_name,
                    "date": pd.Timestamp(d), "pred_change": float(rw_p),
                    "actual_change": float(act), "price_anchor": float(anch),
                })
                records.append({
                    "model": "drift", "horizon": h, "split": split_name,
                    "date": pd.Timestamp(d), "pred_change": float(dr_p),
                    "actual_change": float(act), "price_anchor": float(anch),
                })
    return pd.DataFrame(records)


def _compute_models(
    feat_train:  pd.DataFrame,
    feat_val:    pd.DataFrame,
    feat_test:   pd.DataFrame,
    full_price:  pd.Series,
    baseline_df: pd.DataFrame,
    horizons:    tuple[int, ...] = (1, 7, 30),
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Run tree models + SARIMAX; return (results_df, preds_df)."""
    model_df, model_preds = run_all_models(
        feat_train, feat_val, feat_test, full_price,
        horizons=horizons, baseline_df=baseline_df,
    )
    return model_df, model_preds


def main() -> None:
    print("Loading cleaned splits...")
    train = pd.read_parquet(PROCESSED / "train.parquet")
    val   = pd.read_parquet(PROCESSED / "val.parquet")
    test  = pd.read_parquet(PROCESSED / "test.parquet")

    print("Building feature matrices...")
    feat_train, feat_val, feat_test = build_features(train, val, test)
    save_features(feat_train, feat_val, feat_test, PROCESSED)
    print(f"  feat_train: {feat_train.shape}")
    print(f"  feat_val:   {feat_val.shape}")
    print(f"  feat_test:  {feat_test.shape}")

    from src.evaluate import META_COLS
    feat_cols = [c for c in feat_train.columns if c not in META_COLS]
    feat_stats = {
        "n_features": len(feat_cols),
        "shapes": {
            "train": str(feat_train.shape),
            "val":   str(feat_val.shape),
            "test":  str(feat_test.shape),
        },
    }

    print("Computing staleness statistics...")
    stale_stats = compute_staleness_stats(train, val, test)
    for split_name in ("train", "val", "test"):
        s = stale_stats[split_name]
        print(f"  {split_name:5s}: {s['pct_zero']:.1f}% zero-change  "
              f"({s['n_nonzero']} move-days)  max_run={s['max_run']}")

    print("Computing baselines...")
    full_price = (
        pd.concat([train, val, test])
        .sort_values("Date")
        .set_index("Date")[TARGET_COL]
    )
    baseline_df = _compute_baselines(feat_val, feat_test, full_price)
    print(baseline_df.to_string(index=False))

    print("Running ML models (RF, XGB, LGBM, SARIMAX)...")
    model_df, model_preds_df = _compute_models(
        feat_train, feat_val, feat_test, full_price, baseline_df
    )

    print("Running DL models (LSTM, GRU)...")
    dl_df, dl_preds_df = run_dl_models(
        feat_train, feat_val, feat_test,
        horizons=(1, 7, 30), baseline_df=baseline_df,
    )

    from src.evaluate import build_results_table as _brt
    consolidated_df = _brt(
        pd.concat([baseline_df, model_df, dl_df], ignore_index=True).to_dict("records")
    )
    print(consolidated_df.to_string(index=False))

    consolidated_df.to_parquet(PROCESSED / "model_results.parquet", index=False)

    # ── Collect and save raw predictions ─────────────────────────────────────
    print("Collecting raw predictions for significance tests...")
    baseline_preds_df = _compute_baseline_preds(feat_val, feat_test, full_price)
    all_preds_df = pd.concat(
        [baseline_preds_df, model_preds_df, dl_preds_df], ignore_index=True
    )
    all_preds_df.to_parquet(PROCESSED / "predictions_long.parquet", index=False)
    print(f"  Saved predictions: {len(all_preds_df)} rows × {len(all_preds_df.columns)} cols")

    # ── Diebold-Mariano significance tests ───────────────────────────────────
    print("Running Diebold-Mariano tests...")
    dm_table = run_dm_tests(all_preds_df)
    print(dm_table.to_string(index=False))

    # ── Directional accuracy on move-days ────────────────────────────────────
    dir_acc_table = directional_accuracy_move_days(all_preds_df, h=1, split="test")
    print("\nDir acc (move-days, h=1, test):")
    print(dir_acc_table.to_string(index=False))

    # ── Determine best model at h=1/test ─────────────────────────────────────
    h1_test = consolidated_df[
        (consolidated_df["horizon"] == 1) &
        (consolidated_df["split"] == "test") &
        (~consolidated_df["model"].isin(["random_walk"]))
    ]
    if not h1_test.empty:
        best_model_h1 = h1_test.loc[h1_test["RMSE_skill_%"].idxmax(), "model"]
    else:
        best_model_h1 = "drift"
    print(f"  Best model h=1/test: {best_model_h1}")

    # ── Figures ───────────────────────────────────────────────────────────────
    print("Generating EDA figures...")
    _ensure_figures(train, val, test, feat_train, stale_stats)
    plot_skill_by_horizon(consolidated_df, FIGURES / "fig_skill_by_horizon.png")

    print("Generating significance figures...")
    plot_actual_vs_pred(all_preds_df, best_model_h1, FIGURES / "fig_actual_vs_pred.png")
    plot_error_by_regime(all_preds_df, FIGURES / "fig_error_by_regime.png")

    print("Running SHAP explainability (RF, h=1)...")
    explain_result = run_explain(feat_train, feat_val, feat_test, figures_dir=FIGURES)
    top8_shap = explain_result["top8"]

    print("Computing stationarity statistics (train)...")
    stats = compute_stationarity_tests(train)

    # ── Build Document ────────────────────────────────────────────────────────
    print("Building document...")
    doc = Document()

    title = doc.add_heading("Forecasting Australian Carbon Credit Prices", level=0)
    sub = doc.add_paragraph("A Time-Series Study")
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(13)

    date_para = doc.add_paragraph(
        f"Report generated: {pd.Timestamp.now().strftime('%d %B %Y')}"
    )
    date_para.runs[0].font.size = Pt(10)
    doc.add_page_break()

    doc.add_heading("Contents", level=1)
    for line in [
        "1. Introduction (with Executive Summary)",
        "2. Data and Cleaning",
        "3. Exploratory Analysis",
        "4. Feature Engineering",
        "5. Baselines and Evaluation Framework",
        "6. Modelling Results (RF, XGB, LGBM, SARIMAX)",
        "7. Deep Learning Results (LSTM, GRU)",
        "8. Statistical Significance (Diebold-Mariano Tests)",
        "9. Model Explainability (SHAP)",
        "10. Limitations",
        "11. Conclusion and Future Work",
    ]:
        doc.add_paragraph(line)
    doc.add_page_break()

    build_intro(doc, figures_dir=FIGURES, stats=stats)
    doc.add_page_break()

    build_data_section(doc, figures_dir=FIGURES, stats=stats)
    doc.add_page_break()

    build_eda_section(doc, figures_dir=FIGURES, stats=stats, stale_stats=stale_stats)
    doc.add_page_break()

    build_features_section(doc, figures_dir=FIGURES, feat_stats=feat_stats)
    doc.add_page_break()

    build_baselines_section(doc, baseline_df=baseline_df)
    doc.add_page_break()

    build_modeling_section(doc, figures_dir=FIGURES, consolidated_df=consolidated_df)
    doc.add_page_break()

    build_dl_section(doc, figures_dir=FIGURES, consolidated_df=consolidated_df)
    doc.add_page_break()

    build_significance_section(
        doc, figures_dir=FIGURES, dm_table=dm_table, dir_acc_table=dir_acc_table
    )
    doc.add_page_break()

    build_explainability_section(doc, figures_dir=FIGURES, top8=top8_shap)
    doc.add_page_break()

    build_limitations_section(doc)
    doc.add_page_break()

    build_conclusion_section(doc)

    doc.save(DOCX_OUT)
    print(f"Saved:  {DOCX_OUT}")

    try:
        from docx2pdf import convert
        convert(str(DOCX_OUT), str(PDF_OUT))
        print(f"Saved:  {PDF_OUT}")
    except Exception as exc:
        print(f"Note:   PDF conversion failed ({type(exc).__name__}: {exc})")
        print("        The .docx report is complete. "
              "Convert manually via File -> Export in Word if needed.")


if __name__ == "__main__":
    main()

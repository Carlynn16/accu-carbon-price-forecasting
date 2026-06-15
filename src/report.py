"""
Incremental report builder.

Each public function writes one section into the provided python-docx Document.
Call them in order from scripts/build_report.py.

Sections:
  build_intro(doc, figures_dir, stats)
  build_data_section(doc, figures_dir, stats)
  build_eda_section(doc, figures_dir, stats, stale_stats)
  build_features_section(doc, figures_dir, feat_stats)
  build_baselines_section(doc, baseline_df)
  build_modeling_section(doc, figures_dir, consolidated_df)
  build_dl_section(doc, figures_dir, consolidated_df)
  build_significance_section(doc, figures_dir, dm_table, dir_acc_table)
  build_explainability_section(doc, figures_dir, top8)
  build_limitations_section(doc)
  build_conclusion_section(doc)
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bold_inline(doc: Document, label: str, rest: str) -> None:
    """Add a paragraph with a bold label followed by normal text."""
    para = doc.add_paragraph()
    run = para.add_run(label)
    run.bold = True
    para.add_run(rest)


def _style_table(
    table,
    col_widths_inches: list,
    font_pt: float = 8.5,
) -> None:
    """Apply fixed column widths and compact font size to every cell in a table.

    Forces 'fixed' layout so Word honours the specified widths and numeric
    values never wrap mid-digit across lines.
    """
    # Switch to fixed layout
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    for row in table.rows:
        for j, cell in enumerate(row.cells):
            # Set preferred cell width (dxa = twentieths of a point; 1 inch = 1440 dxa)
            if j < len(col_widths_inches):
                tcPr = cell._tc.get_or_add_tcPr()
                for el in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(el)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(col_widths_inches[j] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

            # Compact spacing + font size
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                for run in para.runs:
                    run.font.size = Pt(font_pt)


def _figure(
    doc: Document,
    img_path: Path,
    caption: str,
    width: float = 5.8,
) -> None:
    """Embed an image (centered) with an italic caption below it."""
    if not img_path.exists():
        doc.add_paragraph(f"[Figure not found: {img_path.name}]")
        return

    doc.add_picture(str(img_path), width=Inches(width))
    # center the picture paragraph (last paragraph added)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────────
# Section 1 — Introduction
# ─────────────────────────────────────────────────────────────────────────────

def build_intro(
    doc: Document,
    figures_dir: Path | None = None,
    stats: dict | None = None,
) -> None:
    """Write Section 1: Introduction."""
    _h(doc, "1. Introduction")

    _p(doc,
       "This report presents a time-series forecasting study of the Australian Generic "
       "carbon-credit spot price (ACCU Generic) at horizons of 1, 7, and 30 calendar days "
       "ahead. Australian Carbon Credit Units (ACCUs) represent one tonne of CO₂-equivalent "
       "abated or sequestered under government-registered emissions reduction methods. The "
       "ACCU Generic price is the benchmark market price, reflecting arms-length transactions "
       "in the secondary market across all eligible methods."
       )

    _p(doc,
       "The guiding methodological principle of this study is that forecasting skill can only "
       "be claimed relative to a naive baseline. The benchmark throughout is the random-walk "
       "model: tomorrow’s forecast equals today’s observed price. A model that fits "
       "well by conventional metrics (R², in-sample RMSE) but cannot beat this baseline "
       "provides no actionable information about future price direction. All performance "
       "comparisons are therefore expressed as percentage improvement in RMSE and MAE over the "
       "random-walk forecast; statistical significance is assessed using the "
       "Diebold–Mariano test."
       )

    _p(doc,
       "A secondary theme is methodological hygiene in feature construction. The raw data "
       "contain columns that are arithmetically derived from the target (lagged changes, "
       "year-to-date percentage returns) and columns that can reconstruct the target level "
       "(sibling-method prices and their premiums over Generic). Using such columns as "
       "predictors produces inflated in-sample accuracy but no genuine forecasting skill. "
       "Section 4 documents the feature audit that excludes these columns before any model "
       "is fitted."
       )

    # ── Executive Summary ─────────────────────────────────────────────────────
    _h(doc, "Executive Summary", level=2)

    _p(doc,
       "Across all three forecast horizons (h = 1, 7, and 30 calendar days) and all six "
       "candidate models — Random Forest, XGBoost, LightGBM, SARIMAX, LSTM, and GRU — "
       "no model achieves a statistically significant improvement in forecast accuracy "
       "over a naive random-walk benchmark on the held-out test set. This result is "
       "confirmed by the Diebold-Mariano test with Harvey-Leybourne-Newbold small-sample "
       "correction (all p-values > 0.05 for the 'better than RW' hypothesis). "
       "Several models are significantly worse than the random walk at h = 7 days."
       )

    _p(doc,
       "The best point estimate is a modest +1.0% RMSE skill score at h = 1 (LSTM, "
       "test set), but this is within the range of statistical noise given the 247-row "
       "test series. A faint directional signal (~57–62% accuracy on genuine-move days "
       "at h = 1, ranging from RF at 56.7% to SARIMAX at 61.9%) is detectable but does "
       "not translate into a meaningful reduction in forecast error."
       )

    _p(doc,
       "SHAP analysis of the best tree model (Random Forest, h = 1) identifies the "
       "dominant drivers as recent realised momentum (chg_0, the most-recent daily change) "
       "and the HIR cross-market price change (hir_chg), followed by chg_1 — a lagged "
       "change that partly reflects the forward-fill artefact — and volume/volatility "
       "metrics. No staleness feature (price_moved, days_since_last_move, moves_7d/30d) "
       "appears in the top 8. These signals are real but weak: they are insufficient to "
       "consistently generate lower forecast error than simply carrying today's price forward."
       )

    _p(doc,
       "The contribution of this study is therefore methodological: it demonstrates "
       "the importance of rigorous benchmark comparisons, the Diebold-Mariano test as "
       "the correct arbiter of forecasting skill, leakage-free feature construction, "
       "and an honest engagement with the near-efficient, highly illiquid structure of "
       "the Australian ACCU spot market. The random walk is the rational benchmark here, "
       "and it is not beaten — a credible and defensible result."
       )


# ─────────────────────────────────────────────────────────────────────────────
# Section 2 — Data and Cleaning
# ─────────────────────────────────────────────────────────────────────────────

def build_data_section(
    doc: Document,
    figures_dir: Path,
    stats: dict | None = None,
) -> None:
    """Write Section 2: Data and Cleaning."""
    _h(doc, "2. Data and Cleaning")

    _p(doc,
       "The dataset consists of daily price and volume observations from the Australian carbon "
       "and renewable-certificate markets, spanning 2 January 2018 to 15 November 2024 "
       "(1,651 trading days). The raw file contained 97 columns covering spot prices across "
       "multiple abatement methods, forward-contract price series, volume figures, "
       "week-on-week and year-to-date change columns, and market-assessment type indicators."
       )

    _h(doc, "2.1  Cleaning Pipeline", level=2)

    _p(doc,
       "A principled, leakage-free cleaning pipeline was applied in the following order to "
       "produce three non-overlapping chronological splits:"
       )

    _bullet(doc,
            "Symbol stripping: dollar signs, percentage signs, and thousands-comma separators "
            "were removed from all value columns before numeric coercion."
            )
    _bullet(doc,
            "High-missingness drop: 41 columns with more than 70% missing values were removed. "
            "These were predominantly forward-contract series (CAL 21–26 and "
            "method-specific analogues) that are only sparsely quoted, plus illiquid-method "
            "columns (No-AD, Landfill Gas, Swaps, Assessment Types)."
            )
    _bullet(doc,
            "Near-constant drop: 6 additional columns where more than 99% of training-set "
            "values were identical were dropped (four calendar-year traded-volume columns and "
            "two CAL 26 change columns)."
            )
    _bullet(doc,
            "Forward-fill imputation on the full chronologically sorted series. This is "
            "strictly past-only: each missing entry is replaced by the most recent prior "
            "observation. Backward-fill was never applied. Carrying the last known price "
            "across a weekend or public holiday is the correct market convention and introduces "
            "no look-ahead."
            )
    _bullet(doc,
            "Residual NaN fill: the small number of leading NaN values at the start of each "
            "series (unreachable by forward-fill) were filled with the training-set column "
            "mean or mode. Validation and test statistics were never consulted."
            )

    _p(doc,
       "After cleaning, 50 columns remained across 1,651 observations with zero missing values."
       )

    _h(doc, "2.2  Chronological Split and Regime Shift", level=2)

    _p(doc,
       "The data were split chronologically by row position: training (70%, 1,155 rows, "
       "2018-01-02 to 2022-12-01), validation (15%, 248 rows, 2022-12-02 to 2023-11-23), "
       "and test (15%, 248 rows, 2023-11-24 to 2024-11-15). No shuffling was performed at "
       "any stage."
       )

    _p(doc,
       "A critical structural feature is the pronounced regime shift between the training "
       "period and the out-of-sample sets. During training, the ACCU Generic price ranged "
       "from A$13.25 to A$57.15 (mean A$21.55), capturing the 2021–22 price spike "
       "driven by policy expectations. The validation and test sets occupy a narrower, "
       "higher plateau (means A$33.66 and A$34.82 respectively, range approximately "
       "A$26–$42). This regime shift has three practical consequences for the analysis:"
       )

    _bullet(doc,
            "It motivates the random-walk as the primary benchmark: a model trained on the "
            "training mean would systematically underpredict the out-of-sample level, so "
            "level forecasting accuracy must be compared against the simplest possible "
            "extrapolation of the most recent observation."
            )
    _bullet(doc,
            "It motivates modelling first differences (daily changes) rather than price "
            "levels, since the distribution of daily changes is far more regime-stable "
            "than the distribution of levels."
            )
    _bullet(doc,
            "It validates the strict chronological split: any reshuffling of rows would "
            "leak post-spike observations into the training set, artificially improving "
            "in-sample fit and producing an over-optimistic evaluation."
            )

    _figure(
        doc,
        figures_dir / "fig_price_timeline.png",
        "Figure 1.  ACCU Generic price over the full sample period (2018-2024). "
        "Shaded regions indicate the training (blue), validation (orange), and test (green) "
        "splits. The 2021-22 price spike is annotated.",
    )
    _figure(
        doc,
        figures_dir / "fig_target_dist_by_split.png",
        "Figure 2.  Price distribution by split (violin plot, left; overlaid histogram, "
        "right). The training distribution is wide and right-skewed due to the 2021-22 price "
        "spike pulling the upper tail; validation and test occupy a narrower high-price regime.",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Section 3 — Exploratory Analysis
# ─────────────────────────────────────────────────────────────────────────────

def build_eda_section(
    doc: Document,
    figures_dir: Path,
    stats: dict,
    stale_stats: dict | None = None,
) -> None:
    """Write Section 3: Exploratory Analysis."""
    _h(doc, "3. Exploratory Analysis")

    _p(doc,
       "All statistical tests and parameter choices reported in this section were computed "
       "on the training split only. Figures show all three splits for visual context."
       )

    # ── 3.1 Stationarity ────────────────────────────────────────────────────
    _h(doc, "3.1  Stationarity Tests", level=2)

    al = stats["adf_level"]
    ac = stats["adf_change"]
    kl = stats["kpss_level"]
    kc = stats["kpss_change"]

    adf_l_p = f"{al['pvalue']:.4f}" if al["pvalue"] > 1e-4 else "< 0.0001"

    _p(doc,
       "The Augmented Dickey–Fuller (ADF) and "
       "Kwiatkowski–Phillips–Schmidt–Shin (KPSS) tests were applied to the "
       "price level and the daily first difference on the training set."
       )

    _bold_inline(
        doc,
        "ADF — price level:  ",
        f"statistic = {al['stat']:.4f}, p-value = {adf_l_p}. "
        f"The null hypothesis of a unit root cannot be rejected (p > 0.05). "
        f"The price level is non-stationary.",
    )
    _bold_inline(
        doc,
        "ADF — daily change:  ",
        f"statistic = {ac['stat']:.4f}, p-value < 0.0001. "
        f"The null hypothesis is strongly rejected. The daily change is stationary.",
    )
    _bold_inline(
        doc,
        "KPSS — price level:  ",
        f"statistic = {kl['stat']:.4f}, p-value {kl['pvalue_str']} (table-bounded). "
        f"H₀ of stationarity is rejected at the 1% level. "
        f"Confirms the level is non-stationary.",
    )
    _bold_inline(
        doc,
        "KPSS — daily change:  ",
        f"statistic = {kc['stat']:.4f}, p-value {kc['pvalue_str']} (table-bounded). "
        f"H₀ cannot be rejected. The daily change is stationary.",
    )

    _p(doc,
       "The joint verdict is unambiguous: the price level is integrated of order one I(1), "
       "and the daily first difference is I(0). Forecasting the level is therefore equivalent "
       "to cumulating a stationary series; the random-walk baseline (which does exactly this) "
       "is hard to beat."
       )

    # ── 3.2 Market Staleness ─────────────────────────────────────────────────
    _h(doc, "3.2  Market Staleness", level=2)

    # Pull staleness numbers — fall back to narrative defaults if not provided
    if stale_stats:
        tr  = stale_stats["train"]
        vl  = stale_stats["val"]
        te  = stale_stats["test"]
        pz_tr, pz_vl, pz_te = tr["pct_zero"], vl["pct_zero"], te["pct_zero"]
        mr_tr = tr["max_run"]
        nz_tr = tr["n_nonzero"]
        # n_valid = len(split) - 1 (first diff is NaN); total rows = n_valid + 1
        total_tr = tr.get("n_valid", nz_tr + round(nz_tr * pz_tr / (100 - pz_tr))) + 1
    else:
        pz_tr, pz_vl, pz_te = 75.1, 59.1, 45.7
        mr_tr, nz_tr = 34, 287
        total_tr = 1155

    _p(doc,
       f"A critical but initially non-obvious feature of this market is its extreme illiquidity. "
       f"The data source records one row per calendar day, but genuine price-discovery events "
       f"— days on which at least one trade executes — are rare. On {pz_tr:.1f}% of training-set "
       f"days, the price does not change at all from the previous observation. The corresponding "
       f"figures are {pz_vl:.1f}% for the validation set and {pz_te:.1f}% for the test set. "
       f"Only {nz_tr} of {total_tr:,} training observations represent genuine price moves."
       )

    _p(doc,
       f"These zero-change days are not independent observations of a flat market; they are "
       f"calendar days on which no trade occurred and the dataset carries forward the "
       f"previous closing price. Stale runs as long as {mr_tr} consecutive days are present "
       f"in the training set (Figure 6). Treating all calendar-day observations as equally "
       f"informative inflates the effective sample size and biases any autocorrelation "
       f"analysis."
       )

    _figure(
        doc,
        figures_dir / "fig_staleness.png",
        "Figure 6.  Left: percentage of zero-change days by split (with number of genuine "
        "move-days and longest stale run annotated). Right: distribution of stale run "
        "lengths in the training set — runs of 5 or more consecutive stale days are "
        "common, with a maximum run exceeding 30 days.",
    )

    # ── 3.3 Autocorrelation — Corrected Analysis ─────────────────────────────
    _h(doc, "3.3  Autocorrelation Structure (Corrected)", level=2)

    lag1_level = stats["acf_level_lags"][0]
    if stale_stats:
        acf_full_lag2 = tr["acf_full_lags"][1]   # lag-2 of full series
        acf_nz_lag1   = tr["acf_nz_lags"][0]     # lag-1 on move-days
        acf_nz_lag2   = tr["acf_nz_lags"][1]     # lag-2 on move-days
    else:
        acf_full_lag2, acf_nz_lag1, acf_nz_lag2 = 0.334, 0.41, 0.142

    _p(doc,
       f"Figure 5 shows the ACF and PACF for both the price level and the daily first "
       f"difference (40 lags, training set only). For the price level, the ACF coefficient "
       f"at lag 1 is {lag1_level:.4f} and remains near 1.0 across all displayed lags, "
       f"consistent with near-random-walk dynamics."
       )

    _p(doc,
       f"The naive ACF of the daily change series shows a prominent spike at lag 2 "
       f"(ACF ≈ {acf_full_lag2:.3f}). This is an artefact of market staleness, not a genuine "
       f"market signal. Because the majority of consecutive observations are identical "
       f"(carried-forward prices), differencing two adjacent stale rows produces two "
       f"consecutive zeros, inflating the lag-2 autocorrelation through the accumulation "
       f"of tied zero-change pairs."
       )

    _p(doc,
       f"Figure 7 isolates the genuine-price-discovery signal by re-computing the ACF "
       f"using only the {nz_tr} non-zero-change training days. On these days, the lag-2 "
       f"coefficient collapses to {acf_nz_lag2:.3f} (confirming it was an artefact), "
       f"while the lag-1 coefficient is {acf_nz_lag1:.3f} — a genuine short-run "
       f"momentum effect: price moves tend to continue in the same direction for one "
       f"additional trading day. This lag-1 momentum signal is the primary linear "
       f"predictive structure exploited in the feature set."
       )

    _figure(
        doc,
        figures_dir / "fig_acf_pacf.png",
        "Figure 5.  ACF and PACF for the price level (top row) and daily change "
        "(bottom row), computed on the training set only (40 lags, 5% confidence bands). "
        "Level ACF ≈ 1.0 across all lags (non-stationary); change ACF shows a spurious "
        "lag-2 spike that disappears on genuine-move days (see Figure 7).",
        width=6.0,
    )
    _figure(
        doc,
        figures_dir / "fig_momentum.png",
        "Figure 7.  ACF of daily price change: full series (blue) vs genuine-move "
        "days only (red), training set. The lag-2 spike collapses from "
        f"≈{acf_full_lag2:.2f} to ≈{acf_nz_lag2:.2f} — a forward-fill artefact. "
        f"The lag-1 momentum signal (≈{acf_nz_lag1:.2f}) on move-days is genuine.",
    )

    # ── 3.4 Returns and Volatility ───────────────────────────────────────────
    _h(doc, "3.4  Returns Distribution and Volatility Clustering", level=2)

    ch = stats["change_stats"]
    _p(doc,
       f"Daily price changes on the training set have mean {ch['mean']:.4f} A$/tonne "
       f"(effectively zero), standard deviation {ch['std']:.4f}, skewness {ch['skew']:.2f}, "
       f"and excess kurtosis {ch['kurtosis']:.0f}. The extreme kurtosis indicates heavy tails: "
       f"large positive or negative moves occur far more frequently than a Gaussian would "
       f"predict. The strong negative skewness reflects occasional large downward repricing "
       f"events. Both properties imply that RMSE and MAE on a few large-move days can "
       f"dominate overall evaluation metrics."
       )

    _p(doc,
       "The rolling 30-day volatility (Figure 4) shows clear volatility clustering. "
       "The 2018–2020 period is relatively calm; a high-volatility episode accompanies "
       "the 2021–22 price spike; and the post-spike period (val and test) is again "
       "calmer. This heteroskedasticity is relevant when interpreting the test-set results: "
       "the evaluation period is a low-volatility regime relative to the most turbulent part "
       "of the training set."
       )

    _figure(
        doc,
        figures_dir / "fig_returns.png",
        "Figure 3.  Daily price changes over the full sample period (left panel) and "
        "the distribution of daily changes on the training set (right panel). "
        "Near-zero mean; heavy tails; extreme excess kurtosis.",
    )
    _figure(
        doc,
        figures_dir / "fig_volatility.png",
        "Figure 4.  Rolling 30-day standard deviation of daily price changes. "
        "Volatility clustering is visible: the 2021-22 spike period shows markedly "
        "higher volatility than the surrounding years.",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Section 4 — Feature Engineering
# ─────────────────────────────────────────────────────────────────────────────

def build_features_section(
    doc: Document,
    figures_dir: Path,
    feat_stats: dict | None = None,
) -> None:
    """Write Section 4: Feature Engineering."""
    _h(doc, "4. Feature Engineering")

    _p(doc,
       "All features are constructed from the full chronologically sorted series using "
       "only lagged values and strictly trailing rolling windows. No centered windows, "
       "no backward-fill, and no information from validation or test rows enters any "
       "feature calculation. The first 30 rows of the training set (the warm-up period "
       "required to populate the 30-day rolling windows) are dropped after feature "
       "construction."
       )

    # ── 4.1 Targets ──────────────────────────────────────────────────────────
    _h(doc, "4.1  Forecast Targets", level=2)

    _p(doc,
       "Three forecast targets are defined, one per horizon h ∈ {1, 7, 30}:"
       )
    _bullet(doc,
            "target_1  =  price(t+1) − price(t)  (1-day-ahead price change)")
    _bullet(doc,
            "target_7  =  price(t+7) − price(t)  (7-day-ahead cumulative change)")
    _bullet(doc,
            "target_30 =  price(t+30) − price(t) (30-day-ahead cumulative change)")

    _p(doc,
       "Modelling the change, not the level, is motivated by the stationarity results in "
       "Section 3.1 (the level is I(1)) and by the pronounced regime shift between the "
       "training and out-of-sample sets. All evaluation metrics are computed on the "
       "change, and the random-walk forecast corresponds to predicting zero change."
       )

    # ── 4.2 Feature Audit — Exclusions ───────────────────────────────────────
    _h(doc, "4.2  Feature Audit: Excluded Columns", level=2)

    _p(doc,
       "Before constructing features, a formal audit excluded three categories of columns "
       "that would introduce data leakage or level-reconstruction artefacts:"
       )

    _bold_inline(doc,
        "Category 1 — Target-derived columns:  ",
        "The raw dataset contains columns that are arithmetic functions of the target "
        "price level: the daily change ($ Change), week-on-week change, YTD percentage "
        "return, and the 7/30/50/100-day SMAs and percentage changes. Retaining any of "
        "these would allow a model to invert them to obtain the contemporaneous price "
        "level, defeating the purpose of modelling changes. All 18 such columns are "
        "excluded."
    )
    _bold_inline(doc,
        "Category 2 — Sibling price levels and premiums:  ",
        "The HIR and SFM ACCU sub-method prices are correlated with the Generic price "
        "but can also reconstruct it: Generic = HIR_price − HIR_premium. Retaining "
        "the sibling price levels or their premium-over-Generic columns would enable "
        "level reconstruction and produce inflated in-sample R² with no genuine "
        "forecasting value. Four price levels and four premium columns are excluded."
    )
    _bold_inline(doc,
        "Category 3 — Raw sibling change columns:  ",
        "The dataset also pre-computes dollar changes and week-on-week changes for HIR "
        "and SFM. These are re-derived from scratch (from the sibling price diffs) to "
        "ensure exact replication of the construction logic. The pre-computed versions "
        "are excluded to avoid double-counting."
    )

    # ── 4.3 Feature Groups ───────────────────────────────────────────────────
    _h(doc, "4.3  Feature Groups", level=2)

    _p(doc,
       "After exclusions, the following feature groups are constructed:"
       )

    _bold_inline(doc,
        "Group A — Momentum (change family, 5 features):  ",
        "chg_0, chg_1, chg_2, chg_3, chg_5. "
        "chg_0 = price(t) − price(t−1): the most-recent realised change, fully "
        "available at the close of day t and the direct carrier of the lag-1 momentum "
        "signal (corr ≈ +0.045 with target_1 on the full series; ≈ +0.41 on genuine-move "
        "days only). chg_1 through chg_5 are further lags. Note: chg_2 shows an "
        "apparent train correlation of ≈ +0.33 — this is largely the forward-fill "
        "artefact identified in Section 3.3 (lag-2 ACF of the full change series), "
        "not a reliable signal; models should not be expected to rely on it out-of-sample."
    )
    _bold_inline(doc,
        "Group B — Staleness features (4 features):  ",
        "price_moved (1/0 flag), days_since_last_move, moves_7d (number of move-days "
        "in trailing 7 days), moves_30d (move-days in trailing 30 days). "
        "These encode the market-liquidity context at each row — essential given that "
        "75% of observations are stale carries."
    )
    _bold_inline(doc,
        "Group C — Volatility regime (2 features):  ",
        "vol_chg_7d and vol_chg_30d: trailing 7- and 30-day rolling standard deviation "
        "of daily changes. These capture local heteroskedasticity."
    )
    _bold_inline(doc,
        "Group D — Volume (4 features, if present):  ",
        "vol_generic_raw (raw traded volume), vol_generic_log1p (log1p-transformed), "
        "vol_generic_trail7 (7-day trailing average), vol_generic_zero (1/0 zero-volume "
        "day flag). Volume is a direct proxy for trading activity."
    )
    _bold_inline(doc,
        "Group E — Calendar (2 features):  ",
        "dow (day of week, 0=Monday) and month. Carbon-credit trading shows day-of-week "
        "and seasonal patterns due to compliance cycles."
    )
    _bold_inline(doc,
        "Group F — Exogenous diffs (up to 8 features):  ",
        "First differences of sibling prices (lgc_chg, stc_chg, hir_chg, sfm_nc_chg, "
        "sfm_cb_chg, erf_chg) and trailing volume averages for HIR and SFM methods. "
        "Levels and premiums are excluded (Category 2 audit above); only changes are used."
    )

    if feat_stats:
        n_feats = feat_stats.get("n_features", "—")
        shapes  = feat_stats.get("shapes", {})
        _p(doc,
           f"After construction and warm-up removal, the feature matrix contains "
           f"{n_feats} features. Training set shape: {shapes.get('train', '—')}; "
           f"validation: {shapes.get('val', '—')}; test: {shapes.get('test', '—')}."
           )

    _figure(
        doc,
        figures_dir / "fig_feature_target_corr.png",
        "Figure 8.  Pearson correlations of each feature with the 1-day-ahead target "
        "(target_1) on the training set. Positive (red) bars indicate momentum-aligned "
        "features; negative (blue) bars indicate mean-reverting signals. Lagged momentum "
        "changes (chg_1, chg_3, chg_5), volume metrics, and calendar features dominate "
        "the top correlation rankings.",
    )


# ─────────────────────────────────────────────────────────────────────────────
# Section 5 — Baselines and Evaluation Framework
# ─────────────────────────────────────────────────────────────────────────────

def build_baselines_section(
    doc: Document,
    baseline_df: "pd.DataFrame | None" = None,
) -> None:
    """Write Section 5: Baselines and Evaluation Framework."""
    import pandas as pd  # local import — report.py has no top-level pandas dependency

    _h(doc, "5. Baselines and Evaluation Framework")

    _p(doc,
       "This section defines the evaluation framework used throughout the study and "
       "establishes the baseline forecasts against which every subsequent model is judged."
       )

    # ── 5.1 The Random-Walk Benchmark ────────────────────────────────────────
    _h(doc, "5.1  The Random-Walk Benchmark", level=2)

    _p(doc,
       "The primary benchmark is the random-walk (RW) forecast: at time t, the "
       "predicted h-day-ahead price is simply price(t). Equivalently, the predicted "
       "h-step change is zero. This is the correct null hypothesis for any I(1) series "
       "and, in illiquid markets such as the ACCU spot market, it is empirically "
       "difficult to beat even with sophisticated methods."
       )

    _p(doc,
       "A secondary, slightly more informed reference is the drift forecast: the "
       "predicted h-step change equals h times the trailing 30-day mean daily change. "
       "This allows for a persistent trend in the recent past without any exogenous "
       "information. Seasonal-naive forecasting is omitted: the ACCU spot market shows "
       "no stable calendar seasonality — price dynamics are policy-driven and dominated "
       "by the staleness structure documented in Section 3.2."
       )

    # ── 5.2 Metrics ──────────────────────────────────────────────────────────
    _h(doc, "5.2  Evaluation Metrics", level=2)

    _p(doc,
       "For all models and baselines, predictions are expressed as h-step price changes. "
       "To evaluate, the predicted level is reconstructed as price_anchor + predicted_change, "
       "where price_anchor = price(t). Metrics are computed on the reconstructed level:"
       )

    _bullet(doc,
            "RMSE and MAE (A$/tonne). Note: because the anchor cancels in the residual, "
            "RMSE on the reconstructed level is identical to RMSE on the raw change — "
            "both are reported in units of A$/tonne and are directly interpretable."
            )
    _bullet(doc,
            "MAPE (%) on the reconstructed level. This is low for short horizons because "
            "most calendar days are stale (zero change, so pred = actual)."
            )
    _bullet(doc,
            "Directional accuracy (%): the fraction of rows where the predicted "
            "sign of the change matches the actual sign. Computed exclusively on "
            "genuine-move rows (actual change ≠ 0), since direction is "
            "meaningless for stale days where both prediction and actuality are zero."
            )

    _p(doc,
       "The headline metric is the skill score:"
       )

    para = doc.add_paragraph()
    run  = para.add_run(
        "RMSE skill score = 100 × (1 − RMSE_model / RMSE_random_walk)   [%]"
    )
    run.italic = True
    para.paragraph_format.left_indent = Inches(0.5)

    _p(doc,
       "A positive skill score means the model beats the random walk. A negative "
       "skill score means it is worse. Zero means tied. The same formula is applied "
       "to MAE. The skill score is the only metric reported in the executive summary."
       )

    # ── 5.3 Walk-Forward CV (deferred) ───────────────────────────────────────
    _h(doc, "5.3  Walk-Forward Cross-Validation", level=2)

    _p(doc,
       "All ML/DL models in subsequent sections will be tuned using walk-forward "
       "(rolling-origin) cross-validation on the training set, implemented via "
       "scikit-learn's TimeSeriesSplit. Standard k-fold cross-validation is never used: "
       "it would validate on the past using folds from the future, producing optimistic "
       "hyperparameter estimates that do not generalise to the deployment horizon. "
       "This section reports only the two baselines, which require no tuning."
       )

    # ── 5.4 Baseline Results ─────────────────────────────────────────────────
    _h(doc, "5.4  Baseline Performance", level=2)

    _p(doc,
       "Tables 1a and 1b report RMSE, MAE, MAPE, and directional accuracy for the "
       "random-walk and drift baselines on the test set (Table 1a) and validation set "
       "(Table 1b). Skill scores for the drift are relative to the random walk. "
       "By construction, the random-walk skill score is 0% for every horizon."
       )

    # Column widths: Model 0.90, h 0.28, RMSE 0.68, MAE 0.68, MAPE% 0.65,
    #                dir_acc% 0.72, RMSE_skill% 0.72  (total ≈ 4.63")
    _T1_WIDTHS = [0.90, 0.28, 0.68, 0.68, 0.65, 0.72, 0.72]
    _T1_COLS   = ["Model", "h", "RMSE", "MAE", "MAPE%", "dir_acc%", "RMSE_skill%"]

    def _build_split_table(split_name: str, label: str) -> None:
        """Build one baseline table for a single split."""
        rows_display = []
        for (model, horizon), grp in baseline_df.groupby(["model", "horizon"]):
            r = grp[grp["split"] == split_name]
            if r.empty:
                continue
            r = r.iloc[0]
            rows_display.append({
                "Model":        model.replace("_", " ").title(),
                "h":            int(horizon),
                "RMSE":         f"{r['RMSE']:.4f}",
                "MAE":          f"{r['MAE']:.4f}",
                "MAPE%":        f"{r['MAPE_%']:.2f}",
                "dir_acc%":     f"{r['dir_acc_%']:.1f}",
                "RMSE_skill%":  f"{r['RMSE_skill_%']:.2f}",
            })
        if not rows_display:
            return
        disp_df = pd.DataFrame(rows_display, columns=_T1_COLS)
        tbl = doc.add_table(rows=len(disp_df) + 1, cols=len(_T1_COLS))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for j, col in enumerate(_T1_COLS):
            hdr[j].text = col
            runs = hdr[j].paragraphs[0].runs
            if runs:
                runs[0].bold = True
        for i, (_, row_vals) in enumerate(disp_df.iterrows()):
            for j, val in enumerate(row_vals):
                tbl.rows[i + 1].cells[j].text = str(val)
        _style_table(tbl, _T1_WIDTHS)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            f"{label}  Baseline performance on the {split_name} set. "
            "RMSE and MAE in A$/tonne. MAPE and dir_acc in %. "
            "Skill scores relative to the random-walk baseline (positive = better)."
        )
        run.italic = True
        run.font.size = Pt(9)

    if baseline_df is not None and len(baseline_df) > 0:
        _p(doc, "")
        _build_split_table("test", "Table 1a.")
        _p(doc, "")
        _build_split_table("val",  "Table 1b.")
    else:
        para = doc.add_paragraph()
        run  = para.add_run("[Baseline results table — to be inserted after running baselines]")
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)


# ─────────────────────────────────────────────────────────────────────────────
# Section 6 — Modelling Results
# ─────────────────────────────────────────────────────────────────────────────

def build_modeling_section(
    doc: Document,
    figures_dir: Path,
    consolidated_df: "pd.DataFrame | None" = None,
) -> None:
    """Write Section 6: ML model results (RF, XGB, LGBM, SARIMAX)."""
    import pandas as pd  # local import

    _h(doc, "6. Modelling Results")

    _p(doc,
       "This section reports the forecasting performance of tree-based ML models "
       "(Random Forest, XGBoost, LightGBM) and a univariate SARIMAX benchmark, "
       "evaluated at horizons h = 1, 7, and 30 days."
       )

    # ── 6.1 Walk-Forward Validation Protocol ─────────────────────────────────
    _h(doc, "6.1  Walk-Forward Validation Protocol", level=2)

    _p(doc,
       "All ML models were tuned using scikit-learn's TimeSeriesSplit with "
       "n_splits = 5 expanding-window folds applied exclusively to the training set. "
       "For each model and horizon the following three-phase protocol was followed:"
       )

    _bullet(doc,
            "Phase 1 — Tune:  GridSearchCV over a predefined hyperparameter grid, "
            "using TimeSeriesSplit(n_splits=5) as the inner cross-validator and "
            "RMSE as the scoring metric. The search space is kept deliberately small "
            "(4–8 candidate configurations) to limit computational cost while still "
            "exploring the most impactful dimensions (tree depth, learning rate, "
            "number of estimators).")
    _bullet(doc,
            "Phase 2 — Val evaluation:  The best estimator from Phase 1 (refit on "
            "the full training set with the winning hyperparameters) is used to "
            "generate predictions for the validation set. No validation data is "
            "seen during Phase 1.")
    _bullet(doc,
            "Phase 3 — Test evaluation:  The best estimator is refit from scratch "
            "on the combined training + validation set (with the same hyperparameters "
            "selected in Phase 1) and used to predict the held-out test set exactly once.")

    _p(doc,
       "Why random k-fold is wrong for time-series data: standard k-fold "
       "cross-validation creates folds by random permutation, so validation examples "
       "from the past are used to assess models trained on the future. This constitutes "
       "look-ahead bias — the hyperparameters selected will appear to generalise "
       "well in-sample but will systematically over-estimate performance on any "
       "future deployment horizon. TimeSeriesSplit enforces strict temporal ordering: "
       "every training fold ends strictly before the corresponding validation fold begins, "
       "replicating the genuine deployment condition."
       )

    # ── 6.2 Model Descriptions ───────────────────────────────────────────────
    _h(doc, "6.2  Model Descriptions", level=2)

    _bold_inline(doc, "Random Forest:  ",
        "Ensemble of decision trees with bootstrap sampling (n_estimators ∈ {100, 300}, "
        "max_depth ∈ {5, unlimited}, min_samples_leaf ∈ {5, 10}). Captures non-linear "
        "interactions between momentum and staleness features.")
    _bold_inline(doc, "XGBoost:  ",
        "Gradient-boosted trees with histogram split finding, L1/L2 regularisation "
        "(n_estimators ∈ {100, 200}, max_depth ∈ {3, 5}, learning_rate ∈ {0.05, 0.1}). "
        "Strong baseline for tabular regression.")
    _bold_inline(doc, "LightGBM:  ",
        "Leaf-wise gradient boosted trees (n_estimators ∈ {100, 200}, "
        "num_leaves ∈ {31, 63}, learning_rate ∈ {0.05, 0.1}). "
        "Faster training and often comparable or superior accuracy to XGBoost.")
    _bold_inline(doc, "SARIMAX(1,1,1):  ",
        "Classical univariate benchmark on the price level with first differencing (d=1). "
        "Fit on the training price series to predict the validation set; refit on "
        "training + validation to predict the test set. Captures linear AR and MA "
        "structure but ignores all exogenous features.")

    # ── 6.3 Consolidated Results ─────────────────────────────────────────────
    _h(doc, "6.3  Consolidated Results", level=2)

    _p(doc,
       "Table 2 shows RMSE and RMSE skill score (relative to the random walk) for all "
       "models on the validation and test sets across the three forecast horizons. "
       "The random-walk skill score is 0% by definition. Positive values indicate "
       "improvement over the random walk; negative values indicate degradation."
       )

    if consolidated_df is not None and len(consolidated_df) > 0:
        _p(doc, "")

        # Build a compact pivot: rows = (model, horizon), cols = val/test metrics
        rows_display = []
        for (model, horizon), grp in consolidated_df.groupby(["model", "horizon"]):
            row = {
                "Model":   model.replace("_", " ").replace("random walk", "Random Walk").title(),
                "h (days)": int(horizon),
            }
            for _, r in grp.iterrows():
                sp = r["split"]
                row[f"{sp} RMSE"]     = f"{r['RMSE']:.4f}"
                row[f"{sp} skill (%)"] = f"{r['RMSE_skill_%']:.2f}"
            rows_display.append(row)

        disp_df = pd.DataFrame(rows_display)

        table = doc.add_table(rows=len(disp_df) + 1, cols=len(disp_df.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(disp_df.columns):
            hdr_cells[j].text = str(col)
            run = hdr_cells[j].paragraphs[0].runs
            if run:
                run[0].bold = True
        for i, (_, row_vals) in enumerate(disp_df.iterrows()):
            data_cells = table.rows[i + 1].cells
            for j, val in enumerate(row_vals):
                data_cells[j].text = str(val)
        # Model 1.20, h(days) 0.48, val RMSE 0.75, val skill 0.80, test RMSE 0.75, test skill 0.80
        _style_table(table, [1.20, 0.48, 0.75, 0.80, 0.75, 0.80])

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            "Table 2.  Consolidated RMSE and skill score (vs random walk) for all models "
            "and baselines on validation and test sets. Skill > 0 means better than random walk."
        )
        run.italic = True
        run.font.size = Pt(9)
    else:
        para = doc.add_paragraph()
        run  = para.add_run(
            "[Model results table — run scripts/run_models.py and regenerate report]"
        )
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── 6.4 Skill by Horizon ─────────────────────────────────────────────────
    _h(doc, "6.4  Skill Score by Horizon", level=2)

    skill_fig = figures_dir / "fig_skill_by_horizon.png"
    _figure(
        doc,
        skill_fig,
        "Figure 9.  RMSE skill score (%) vs random walk by model and horizon. "
        "Left: validation set. Right: test set. Bars above zero indicate improvement "
        "over the random-walk benchmark. Error bars omitted (single evaluation per split).",
    )

    # ── 6.5 Interpretation ───────────────────────────────────────────────────
    _h(doc, "6.5  Interpretation", level=2)

    _p(doc,
       "The results should be interpreted with awareness of the market's structural "
       "characteristics established in earlier sections:"
       )

    _bullet(doc,
            "The test set occupies a low-volatility regime (Section 3.4). Random-walk "
            "RMSE on the test set is materially lower than on validation (h=1: 0.32 vs 0.58 "
            "A$/tonne), so a positive skill score on the test set reflects genuine signal "
            "relative to a weaker baseline, not superior accuracy in absolute terms.")
    _bullet(doc,
            "Market staleness (75.1% stale days in training, ~46% in test) means that a "
            "large fraction of rows have target = 0. Any model that predicts near-zero "
            "changes will achieve a low RMSE simply by being right on stale days; "
            "the directional accuracy metric (Table 1) is more diagnostic for genuine "
            "price-discovery days.")
    _bullet(doc,
            "Tree model performance on the test set: all three tree models (RF, XGB, "
            "LightGBM) show negative or negligible skill at h = 1 and deteriorate "
            "further at h = 7 and h = 30. XGB and LightGBM are significantly worse "
            "than the random walk at h = 7 (positive DM stat, p < 0.05). The 25-feature "
            "matrix captures momentum (chg_0) and staleness structure, but these signals "
            "are insufficient to overcome the near-efficient, event-driven price dynamics.")
    _bullet(doc,
            "SARIMAX(1,1,1) result: close to but not significantly better than the "
            "random walk at h = 1 (skill ≈ +0.5%), and worse at longer horizons. "
            "Because the daily price-change series has minimal linear autocorrelation "
            "(Section 3.3) and SARIMAX ignores all exogenous features, this result is "
            "the expected confirmation that linear AR/MA structure alone carries no "
            "forecasting value here.")


# ─────────────────────────────────────────────────────────────────────────────
# Section 7 — Deep Learning Results (LSTM, GRU)
# ─────────────────────────────────────────────────────────────────────────────

def build_dl_section(
    doc: Document,
    figures_dir: Path,
    consolidated_df: "pd.DataFrame | None" = None,
) -> None:
    """Write Section 7: LSTM and GRU results."""
    import pandas as pd

    _h(doc, "7. Deep Learning Results (LSTM, GRU)")

    _p(doc,
       "This section presents results for two recurrent neural network architectures: "
       "Long Short-Term Memory (LSTM) and Gated Recurrent Unit (GRU). Both are trained "
       "to predict the h-step price change using the same 25-feature matrix as the "
       "tree models, but consume the features as a temporal sequence rather than a "
       "single feature vector. Detailed implementation notes are provided below."
       )

    # ── 7.1 Sequence Construction ─────────────────────────────────────────────
    _h(doc, "7.1  Sequence Construction", level=2)

    _p(doc,
       "Raw features are standardised using a StandardScaler fit on the training set "
       "only (mean and standard deviation computed from training rows; applied to "
       "validation and test). The scaled feature matrix is then cut into sliding "
       "windows of length L = 20 trading days."
       )
    _bullet(doc,
            "Window at anchor row t: input = scaled rows [t-19 .. t] (shape 20 x 25), "
            "target = target_h(t) = price(t+h) - price(t).")
    _bullet(doc,
            "Assignment: each sample belongs to the split of its anchor row t. "
            "A validation window may look back into the training tail — this is "
            "past-only and does not constitute leakage (the scaler is already fit, "
            "and those rows precede t chronologically).")
    _bullet(doc,
            "Warm-up: the first L-1 rows of the series cannot form a complete window "
            "and are excluded. Windows whose target is NaN (last h rows) are also dropped.")

    # ── 7.2 Model Architectures ───────────────────────────────────────────────
    _h(doc, "7.2  Model Architectures", level=2)

    _bold_inline(doc, "LSTM (Long Short-Term Memory):  ",
        "Single-layer LSTM with hidden size 64. At each time step the LSTM updates "
        "a hidden state h_t and a memory cell c_t via learned gating mechanisms "
        "(input, forget, output gates). The final hidden state h_T is passed through "
        "a dropout layer (p = 0.2) and a linear head to produce the scalar prediction.")
    _bold_inline(doc, "GRU (Gated Recurrent Unit):  ",
        "Single-layer GRU with hidden size 64. GRU simplifies LSTM by merging the "
        "cell and hidden state into a single h_t using two gates (reset and update). "
        "Otherwise identical architecture: dropout 0.2, linear head.")

    _p(doc,
       "Both models use the Adam optimiser (lr = 0.001), MSE loss on the predicted "
       "price change, and mini-batch training (batch size 32) without shuffling "
       "(chronological order preserved within each epoch)."
       )

    # ── 7.3 Training Protocol ─────────────────────────────────────────────────
    _h(doc, "7.3  Training Protocol", level=2)

    _p(doc,
       "The protocol mirrors the C2 tree-model protocol exactly:"
       )
    _bullet(doc,
            "Phase 1 — Train on TRAIN sequences. Monitor validation MSE after every "
            "epoch; save the best weights. Stop when validation MSE has not improved "
            "by more than 1e-7 for 10 consecutive epochs (patience = 10), with a "
            "maximum of 100 epochs.")
    _bullet(doc,
            "Phase 3 — Refit on TRAIN+VAL sequences for exactly as many epochs as "
            "Phase 1 selected (using the same random seed). Evaluate once on TEST.")
    _bullet(doc,
            "Fixed seeds (torch, numpy, random = 42) ensure fully reproducible results.")

    # ── 7.4 Results ───────────────────────────────────────────────────────────
    _h(doc, "7.4  Results", level=2)

    _p(doc,
       "Table 3 shows the complete consolidated results including LSTM and GRU "
       "alongside all baselines and C2 models. Skill scores are relative to the "
       "random-walk baseline (0% by definition)."
       )

    if consolidated_df is not None and len(consolidated_df) > 0:
        _p(doc, "")

        rows_display = []
        for (model, horizon), grp in consolidated_df.groupby(["model", "horizon"]):
            row = {
                "Model":    model.replace("_", " ").title(),
                "h (days)": int(horizon),
            }
            for _, r in grp.iterrows():
                sp = r["split"]
                row[f"{sp} RMSE"]     = f"{r['RMSE']:.4f}"
                row[f"{sp} skill (%)"] = f"{r['RMSE_skill_%']:.2f}"
            rows_display.append(row)

        disp_df = pd.DataFrame(rows_display)
        table = doc.add_table(rows=len(disp_df) + 1, cols=len(disp_df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for j, col in enumerate(disp_df.columns):
            hdr[j].text = str(col)
            run = hdr[j].paragraphs[0].runs
            if run:
                run[0].bold = True
        for i, (_, row_vals) in enumerate(disp_df.iterrows()):
            cells = table.rows[i + 1].cells
            for j, val in enumerate(row_vals):
                cells[j].text = str(val)
        # Model 1.20, h(days) 0.48, val RMSE 0.75, val skill 0.80, test RMSE 0.75, test skill 0.80
        _style_table(table, [1.20, 0.48, 0.75, 0.80, 0.75, 0.80])

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            "Table 3.  Complete consolidated results: all baselines, C2 models, and "
            "DL models. RMSE in A$/tonne. Skill score vs random walk (positive = better)."
        )
        run.italic = True
        run.font.size = Pt(9)
    else:
        para = doc.add_paragraph()
        run  = para.add_run(
            "[DL results table — run scripts/build_report.py to populate]"
        )
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── 7.5 Skill by horizon (updated figure) ────────────────────────────────
    _h(doc, "7.5  Skill Score by Horizon (All Models)", level=2)

    _figure(
        doc,
        figures_dir / "fig_skill_by_horizon.png",
        "Figure 9 (updated).  RMSE skill score (%) vs random walk for all models "
        "including LSTM and GRU. Left: validation set. Right: test set. "
        "Bars above zero indicate improvement over the random-walk benchmark.",
    )

    # ── 7.6 Interpretation ────────────────────────────────────────────────────
    _h(doc, "7.6  Why DL Does Not Beat the Random Walk", level=2)

    _p(doc,
       "The deep-learning results confirm the expected finding: neither LSTM nor "
       "GRU achieves positive RMSE skill over the random walk at any horizon. "
       "Three structural reasons explain this:"
       )
    _bullet(doc,
            "Sample scarcity. The TRAIN set contains approximately 1,100 rows, and "
            "after windowing (L = 20) only ~1,080 sequences are available. Recurrent "
            "neural networks are data-hungry; generalisation from so few sequences is "
            "unreliable. The tree models (which see the full feature vector without "
            "windowing) have a modest advantage here.")
    _bullet(doc,
            "Target sparsity. In the training set, 75% of target_h values are zero "
            "(stale-day price unchanged). The MSE loss is dominated by these zero "
            "rows, pushing the network toward predicting near-zero changes for "
            "every origin — which is what the random walk already does.")
    _bullet(doc,
            "Weak temporal signal. The ACF analysis (Section 3.3) showed that "
            "autocorrelation of price changes becomes insignificant within 1-2 lags. "
            "A 20-step lookback window contains mostly noise beyond the immediate lag, "
            "giving the recurrent layers little to learn from.")
    _p(doc,
       "These results are not a failure of the methodology — they are the honest "
       "finding. In an interview context, the ability to explain why a sophisticated "
       "model does not add value, and to back that explanation with diagnostic "
       "evidence, is more valuable than an inflated benchmark score."
       )


# ─────────────────────────────────────────────────────────────────────────────
# Section 8 — Statistical Significance (Diebold-Mariano Tests)
# ─────────────────────────────────────────────────────────────────────────────

def build_significance_section(
    doc:           "Document",
    figures_dir:   Path,
    dm_table:      "pd.DataFrame | None" = None,
    dir_acc_table: "pd.DataFrame | None" = None,
) -> None:
    """Write Section 8: Diebold-Mariano significance tests + error analysis."""
    import pandas as pd

    _h(doc, "8. Statistical Significance (Diebold-Mariano Tests)")

    _p(doc,
       "Sections 6 and 7 reported point estimates of RMSE skill. This section "
       "answers a stricter question: are any of the observed performance gaps "
       "statistically distinguishable from chance, given the length of the test "
       "series and the properties of the forecast errors?"
       )

    # ── 8.1 Test Design ──────────────────────────────────────────────────────
    _h(doc, "8.1  Test Design", level=2)

    _p(doc,
       "The Diebold-Mariano (DM) test (Diebold & Mariano, 1995) compares two "
       "competing forecast sequences by constructing the loss-differential series:"
       )
    _bullet(doc,
            "d_t = L(e_model,t) − L(e_RW,t),  where L is squared error "
            "(e_model,t = y_{t+h} − ŷ_{model,t+h}).")
    _bullet(doc,
            "H₀: E[d_t] = 0 (equal predictive accuracy). "
            "H₁ (two-sided): models differ in expected loss.")
    _bullet(doc,
            "DM statistic: d̄ / sqrt(V̂/T), where V̂ is the Newey-West "
            "(1987) HAC long-run variance with M = h−1 Bartlett lags — "
            "accounting for the serial correlation that arises naturally in "
            "h-step-ahead forecast errors.")
    _bullet(doc,
            "Small-sample correction: Harvey, Leybourne & Newbold (1997) "
            "multiply the DM statistic by c = sqrt((T+1−2h+T⁻¹h(h−1))/T) "
            "and compare against t(T−1) rather than N(0,1). "
            "With T ≈ 220–247 observations this correction is material.")
    _bullet(doc,
            "Significance level: α = 0.05 (two-sided). "
            "Negative DM stat → model outperforms RW. "
            "Positive DM stat → model underperforms RW.")

    # ── 8.2 Results ──────────────────────────────────────────────────────────
    _h(doc, "8.2  Results: No Model Significantly Outperforms the Random Walk", level=2)

    _p(doc,
       "Table 4 shows the DM test results for every model at horizons "
       "h ∈ {1, 7, 30} on the held-out test set (and, for comparison, on the "
       "validation set). The headline finding is unambiguous:"
       )
    _bullet(doc,
            "No model achieves a statistically significant improvement over the "
            "random walk at any horizon on the test set (all p-values > 0.05).")
    _bullet(doc,
            "Small positive test skills at h=1 (LSTM ≈ +1.0%, SARIMAX ≈ +0.5%, "
            "drift ≈ +0.2%) are within the noise range. Their DM statistics are "
            "small and p-values are large.")
    _bullet(doc,
            "Several models are significantly WORSE than the random walk at "
            "h=7 and h=30 (positive DM stat, p < 0.05), confirming that "
            "attempting to forecast longer horizons with these features is "
            "counter-productive.")

    if dm_table is not None and len(dm_table) > 0:
        _p(doc, "")

        # Filter to test only for the main table; include val for context
        test_dm = dm_table[dm_table["split"] == "test"].copy()

        rows_disp: list[dict] = []
        for _, r in test_dm.iterrows():
            rows_disp.append({
                "Model":   str(r["model"]).replace("_", " ").title(),
                "h":       int(r["horizon"]),
                "n_obs":   int(r["n_obs"]),
                "DM stat": f"{r['DM_stat']:+.3f}",
                "p-value": f"{r['p_value']:.4f}",
                "Verdict": str(r["verdict"]),
            })

        disp_df = pd.DataFrame(rows_disp)
        table   = doc.add_table(rows=len(disp_df) + 1, cols=len(disp_df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for j, col in enumerate(disp_df.columns):
            hdr[j].text = str(col)
            run = hdr[j].paragraphs[0].runs
            if run:
                run[0].bold = True
        for i, (_, row_v) in enumerate(disp_df.iterrows()):
            for j, val in enumerate(row_v):
                table.rows[i + 1].cells[j].text = str(val)
        # Model 1.20, h 0.28, n_obs 0.55, DM stat 0.72, p-value 0.72, Verdict 1.60
        _style_table(table, [1.20, 0.28, 0.55, 0.72, 0.72, 1.60])

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            "Table 4.  Diebold-Mariano test results (test set). "
            "Loss = squared error. HAC: Newey-West Bartlett, M = h−1 lags. "
            "HLN small-sample correction applied. α = 0.05, two-sided t(T−1). "
            "Negative DM stat = model has lower loss than random walk."
        )
        run.italic = True
        run.font.size = Pt(9)
    else:
        para = doc.add_paragraph()
        run  = para.add_run("[DM results table — run scripts/build_report.py to populate]")
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    # ── 8.3 Error Analysis by Market Regime ──────────────────────────────────
    _h(doc, "8.3  Error Analysis by Market Regime", level=2)

    _p(doc,
       "ACCU prices exhibit a bimodal behaviour: roughly half the trading days "
       "in the test set record no price change (stale days), while the remainder "
       "are genuine-move days where prices shift. These two regimes have very "
       "different forecasting difficulty."
       )

    _figure(
        doc,
        figures_dir / "fig_actual_vs_pred.png",
        "Figure D1.  Actual next-day price level vs random-walk and best-model "
        "forecasts on the test set (h = 1). The three lines nearly overlap, "
        "illustrating why marginal RMSE differences are statistically insignificant.",
    )

    _figure(
        doc,
        figures_dir / "fig_error_by_regime.png",
        "Figure D2.  RMSE decomposed into stale days (actual price change = 0) "
        "and genuine-move days (actual price change ≠ 0), h = 1, test set. "
        "All models achieve near-zero RMSE on stale days (trivially correct) "
        "but have substantially higher errors on move days, confirming that the "
        "genuine forecasting challenge lies in predicting when and by how much "
        "prices move.",
    )

    if dir_acc_table is not None and len(dir_acc_table) > 0:
        _p(doc,
           "Table 5 shows directional accuracy restricted to genuine-move days "
           "at h = 1 on the test set. The random walk — which always predicts "
           "zero change — has 0% directional accuracy on move days by construction. "
           "Models with positive directional accuracy are predicting the direction "
           "of price moves, even if the magnitude is imprecise."
           )
        _p(doc, "")

        da_disp = dir_acc_table.copy()
        da_disp["model"] = da_disp["model"].str.replace("_", " ").str.title()
        da_disp.columns  = ["Model", "n move-days", "Dir Acc move-days (%)"]

        table2 = doc.add_table(rows=len(da_disp) + 1, cols=len(da_disp.columns))
        table2.style = "Table Grid"
        hdr2 = table2.rows[0].cells
        for j, col in enumerate(da_disp.columns):
            hdr2[j].text = str(col)
            run = hdr2[j].paragraphs[0].runs
            if run:
                run[0].bold = True
        for i, (_, row_v) in enumerate(da_disp.iterrows()):
            for j, val in enumerate(row_v):
                table2.rows[i + 1].cells[j].text = str(val)
        # Model 1.50, n move-days 0.95, Dir Acc move-days (%) 1.30
        _style_table(table2, [1.50, 0.95, 1.30])

        cap2 = doc.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap2.add_run(
            "Table 5.  Directional accuracy on genuine-move days only, h = 1, test set. "
            "Excludes stale days (Δprice = 0) where any near-zero prediction is trivially correct."
        )
        run2.italic = True
        run2.font.size = Pt(9)

    # ── 8.4 Validation vs Test Overfitting Gap ────────────────────────────────
    _h(doc, "8.4  Validation vs Test Overfitting Gap", level=2)

    _p(doc,
       "Comparing the DM test verdicts on the validation and test sets reveals "
       "a recurring pattern: some models that appeared marginally better than "
       "the random walk on validation fail to replicate that edge on the test set. "
       "This is the hallmark of mild overfitting to the validation regime."
       )
    _bullet(doc,
            "The validation period (mid-2021 to late-2022) covered the market's "
            "high-volatility, high-price period. Features capturing momentum and "
            "regime shifts were genuinely informative during this window.")
    _bullet(doc,
            "The test period (late-2022 to late-2024) is a lower-volatility, "
            "range-bound regime. The random-walk RMSE itself is materially lower "
            "(h=1: ~0.32 A$/tonne vs ~0.58 on val), leaving less absolute room "
            "for any model to achieve a meaningfully lower error.")
    _bullet(doc,
            "The DM test on the test set is the appropriate final arbiter. "
            "Val-period skill is informative for understanding model behaviour "
            "but does not constitute out-of-sample evidence.")
    _p(doc,
       "The study's conclusion is therefore conservative and defensible: "
       "no candidate model in this comparison achieves statistically significant "
       "improvement over the random walk for ACCU spot price forecasting at "
       "horizons 1, 7, or 30 days, given the available data."
       )


# ─────────────────────────────────────────────────────────────────────────────
# Section 9 — Model Explainability (SHAP)
# ─────────────────────────────────────────────────────────────────────────────

def build_explainability_section(
    doc:         "Document",
    figures_dir: Path,
    top8:        "pd.DataFrame | None" = None,
) -> None:
    """Write Section 9: SHAP explainability for the best tree model (RF, h=1)."""
    import pandas as pd

    _h(doc, "9. Model Explainability (SHAP)")

    _p(doc,
       "This section explains what the best-performing tree model (Random Forest at h=1) "
       "actually learned from the data. The goal is not to justify the model's performance — "
       "Section 8 established that no model significantly beats the random walk — but to "
       "understand which features drove the model's decisions and whether those patterns "
       "are economically meaningful or artefacts of the data structure."
       )

    # ── 9.1 Method ────────────────────────────────────────────────────────────
    _h(doc, "9.1  Method: SHAP TreeExplainer", level=2)

    _p(doc,
       "SHAP (SHapley Additive exPlanations; Lundberg & Lee, 2017) assigns each feature "
       "a contribution value for each individual prediction. For tree ensembles, "
       "TreeExplainer computes exact SHAP values in polynomial time using the tree "
       "structure — no sampling approximation."
       )
    _bullet(doc,
            "Additivity: for each test observation, SHAP values sum exactly to "
            "model_output(x) − E[f(X)], where E[f(X)] is the base value "
            "(model's mean output over training). Verified by the test suite.")
    _bullet(doc,
            "Feature importance: mean |SHAP| across the test set is used as the "
            "global importance metric — proportional to each feature's average "
            "absolute impact on the predicted price change.")
    _bullet(doc,
            "The same RF hyperparameters and TRAIN+VAL refit procedure as Block C2 "
            "were used, so the SHAP values reflect the model actually evaluated in "
            "Section 6.")

    # ── 9.2 Top Features ──────────────────────────────────────────────────────
    _h(doc, "9.2  Top Features by Mean |SHAP|", level=2)

    _p(doc,
       "Table 6 and Figure E1 show the feature importance ranking. The top features are:"
       )
    _bullet(doc,
            "chg_0 (mean |SHAP| = 0.023 A$/tonne) — the most-recent realised daily "
            "price change. This is the primary momentum carrier identified in the "
            "ACF analysis (Section 3.3): a genuine lag-1 signal on active-trading days.")
    _bullet(doc,
            "hir_chg (0.020) — the first difference of the HIR (Human-Induced "
            "Regeneration) ACCU price. HIR and Generic prices trade in the same "
            "market; co-movement of their daily changes provides a cross-market "
            "confirmation signal.")
    _bullet(doc,
            "chg_1 (0.020) — the lagged change at t−2. This carries a mix of genuine "
            "short-run momentum persistence and the forward-fill artefact discussed "
            "in Section 9.3.")
    _bullet(doc,
            "hir_vol_trail7 (0.012), vol_generic_trail7 (0.008), vol_chg_7d (0.004) — "
            "trailing volume and volatility metrics. These proxy for market activity "
            "and liquidity-regime context.")
    _bullet(doc,
            "chg_5 (0.008) and month (0.004) — a longer-lag momentum feature and a "
            "calendar seasonality signal.")
    _p(doc,
       "Notably, no staleness feature (price_moved, days_since_last_move, moves_7d, "
       "moves_30d) appears in the top 8. The model's dominant signals are momentum "
       "and cross-market co-movement — not regime classification per se. However, "
       "these signals are weak in magnitude (largest mean |SHAP| = 0.023 A$/tonne "
       "vs a random-walk RMSE of 0.32 A$/tonne on the test set) and insufficient "
       "to consistently beat the random-walk baseline."
       )

    if top8 is not None and len(top8) > 0:
        _p(doc, "")
        disp = top8.copy()
        disp.columns = ["Feature", "Mean |SHAP| (A$/tonne)"]

        table = doc.add_table(rows=len(disp) + 1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for j, col in enumerate(disp.columns):
            hdr[j].text = str(col)
            run = hdr[j].paragraphs[0].runs
            if run:
                run[0].bold = True
        for i, (_, row_v) in enumerate(disp.iterrows()):
            for j, val in enumerate(row_v):
                table.rows[i + 1].cells[j].text = str(val)
        # Feature 1.80, Mean |SHAP| (A$/tonne) 1.80
        _style_table(table, [1.80, 1.80])

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(
            "Table 6.  Top-8 features by mean |SHAP| for Random Forest at h = 1 "
            "on the test set. SHAP values measure each feature's average absolute "
            "impact on the predicted price change (A$/tonne)."
        )
        run.italic = True
        run.font.size = Pt(9)

    _figure(
        doc,
        figures_dir / "fig_shap_summary.png",
        "Figure E1.  Feature importance for Random Forest (h = 1, test set) ranked "
        "by mean |SHAP| value. Red bars = momentum lag features (chg_*); "
        "green = staleness features; blue = volatility, volume, exogenous, calendar. "
        "The dominant drivers are chg_0 (recent momentum) and hir_chg (HIR cross-market "
        "change), with volume/volatility metrics in the mid-range. No staleness feature "
        "appears in the top 8.",
    )

    # ── 9.3 chg_1 Dependence Plot ─────────────────────────────────────────────
    _h(doc, "9.3  The Forward-Fill Artefact in chg_1", level=2)

    _p(doc,
       "The dependence plot for chg_1 (the penultimate day's price change) in "
       "Figure E2 reveals an important pattern. Large absolute values of chg_1 "
       "— corresponding to a genuine price move two days ago — are associated with "
       "elevated SHAP magnitude, but the direction of the SHAP contribution is "
       "coloured by chg_0 (yesterday's change)."
       )
    _bullet(doc,
            "When chg_1 is large and positive but chg_0 is zero (a stale day "
            "following a genuine up-move), the model has learned to expect continued "
            "staleness — this is a real and useful regime signal.")
    _bullet(doc,
            "However, the same pattern is partly an artefact of the forward-fill "
            "imputation: after a genuine move at t−2, the imputed values at t−1 and "
            "t retain the same price level (zero change), creating a mechanical "
            "sequence of (non-zero chg_1, zero chg_0) that dominates the feature "
            "space on stale days. The model learns this sequence but it does not "
            "represent a tradeable momentum signal — it is the market being closed "
            "or un-traded, not a directional forecast.")
    _bullet(doc,
            "Practically: the model's attention to chg_1 is approximately equivalent "
            "to detecting 'was there a move two days ago?' — which is a staleness "
            "signal, not a price-direction signal. This is consistent with near-zero "
            "directional accuracy on genuine move-days (Section 8.3).")

    _figure(
        doc,
        figures_dir / "fig_shap_dependence_chg1.png",
        "Figure E2.  SHAP dependence plot for chg_1 (change at t−2), h = 1, test set. "
        "Each point is one test observation; colour = chg_0 (most-recent change). "
        "The pattern shows that large prior moves raise the model's attention — "
        "partly a genuine staleness regime signal, partly the forward-fill artefact "
        "where a move at t−2 is followed mechanically by zero chg_0.",
    )

    # ── 9.4 Why SHAP Confirms the DM Finding ─────────────────────────────────
    _h(doc, "9.4  Interpretation: Regime Detection, Not Price Discovery", level=2)

    _p(doc,
       "Taken together, the SHAP analysis resolves the apparent paradox of Section 6: "
       "why does the Random Forest achieve marginally positive skill (+−3%) at h=1 "
       "on some data splits, but fail to beat the random walk significantly?"
       )
    _bullet(doc,
            "The top SHAP features — chg_0, hir_chg, chg_1 — are all short-lag "
            "momentum and cross-market signals. These carry genuine information "
            "about the direction and magnitude of recent price moves. On stale days "
            "(~46% of test), both the model and the random walk predict near-zero "
            "change; the model offers no advantage here.")
    _bullet(doc,
            "On genuine-move days (where skill would matter most), the model's "
            "directional accuracy is ~57–62% (Section 8.3), marginally above the "
            "50% random threshold. The momentum and cross-market signals are real "
            "but do not produce large enough SHAP magnitudes (max 0.023 A$/tonne) "
            "to dominate the residual error on these high-variance days.")
    _bullet(doc,
            "The implication for longer horizons (h = 7, 30) is direct: short-lag "
            "momentum (chg_0, chg_1) decays within one to two steps, and the "
            "cross-market HIR co-movement also becomes less predictable. At longer "
            "horizons the model reverts toward zero-change prediction — matching "
            "the random walk without consistently beating it.")

    # ── 9.5 DL Explainability Remark ─────────────────────────────────────────
    _h(doc, "9.5  Deep Learning Explainability (Qualitative)", level=2)

    _p(doc,
       "SHAP TreeExplainer is not directly applicable to the recurrent architectures "
       "(LSTM, GRU) trained in Section 7. Gradient-based methods (Integrated Gradients, "
       "GradCAM) or SHAP KernelExplainer (model-agnostic, slow) could be applied, "
       "but given that both DL models fail to beat the random walk at any horizon, "
       "detailed attribution analysis would likely identify a similar pattern to the RF: "
       "short-lag momentum features (chg_0, chg_1) as primary drivers, with the network "
       "reverting to near-zero predictions on stale days where the training loss is "
       "dominated by zero-target rows. This is a limitation acknowledged for future work."
       )


# ─────────────────────────────────────────────────────────────────────────────
# Section 10 — Limitations
# ─────────────────────────────────────────────────────────────────────────────

def build_limitations_section(doc: "Document") -> None:
    """Write Section 10: Limitations."""
    _h(doc, "10. Limitations")

    _p(doc,
       "This study is designed to be honest about the boundaries of its conclusions. "
       "The following limitations should be considered when interpreting the results."
       )

    _h(doc, "10.1  Small Effective Sample", level=2)
    _p(doc,
       "The cleaned training set contains 1,155 rows, but only approximately 287 of "
       "these are genuine price-discovery days (days with non-zero price change). "
       "Statistical learning on 287 informative observations is severely constrained: "
       "tree models cannot reliably learn interactions between more than a handful of "
       "features, and sequence models (LSTM, GRU) face an even more acute sample problem "
       "after windowing. The true signal-to-noise ratio is far lower than the raw row "
       "count suggests."
       )

    _h(doc, "10.2  Single Chronological Split and Regime Shift", level=2)
    _p(doc,
       "All models are evaluated on a single held-out test set covering late 2023 to "
       "late 2024 — a low-volatility, range-bound regime that differs materially from "
       "the high-volatility training period (2018–2022). This single-split evaluation "
       "cannot distinguish between a model that generalises well across market regimes "
       "and one that happens to perform well on this specific test window. Nested "
       "walk-forward cross-validation across multiple non-overlapping test windows would "
       "produce more reliable estimates, but requires a dataset roughly 3–5× larger than "
       "the one available here."
       )
    _bullet(doc,
            "The staleness fraction drops from 75.1% in training to 45.7% in test, "
            "meaning the test set is structurally different from the training set in "
            "exactly the dimension that dominates model behaviour (stale vs active regime)."
            )

    _h(doc, "10.3  Daily Frequency and Aggregation Choice", level=2)
    _p(doc,
       "Forecasting at the daily frequency amplifies the staleness problem: the majority "
       "of calendar days are non-trading days where the price dataset carries forward "
       "the last known value. A weekly or move-day-conditional modelling framework "
       "would substantially improve the effective sample size and the signal-to-noise "
       "ratio, at the cost of losing sub-weekly resolution."
       )

    _h(doc, "10.4  Limited Exogenous Feature Set", level=2)
    _p(doc,
       "The exogenous features used (sibling certificate price changes, HIR/SFM volume "
       "metrics) are all drawn from the same market-data file. Several potentially "
       "important drivers of ACCU price movements are absent:"
       )
    _bullet(doc,
            "EU-ETS (European carbon market) prices and global carbon market sentiment.")
    _bullet(doc,
            "Australian energy prices (electricity, gas) and renewable-energy certificate "
            "supply (LGC issuance, Renewable Energy Target compliance).")
    _bullet(doc,
            "Macro variables: AUD/USD exchange rate, commodity indices, interest rates.")
    _bullet(doc,
            "Policy events: Australian government policy announcements, Safeguard "
            "Mechanism reforms, ERF auction outcomes.")
    _bullet(doc,
            "Order-book and market-microstructure data (bid-ask spread, depth, "
            "OTC versus exchange volumes) that would directly capture liquidity regime.")
    _p(doc,
       "The absence of these variables is the most likely explanation for why no model "
       "consistently beats the random walk: ACCU price moves appear to be primarily "
       "driven by discrete, policy-event-triggered shocks that are not predictable "
       "from lagged market data alone."
       )

    _h(doc, "10.5  Light DL Tuning", level=2)
    _p(doc,
       "The LSTM and GRU hyperparameter space is narrow by design (single layer, fixed "
       "hidden size 64, fixed sequence length 20, dropout 0.2). A thorough architecture "
       "search would include varying hidden size, number of layers, sequence length, "
       "dropout rate, and learning-rate schedule. Given the small training set and the "
       "tree-model results, there is no strong prior reason to expect a more thoroughly "
       "tuned RNN to produce materially different conclusions — but it cannot be ruled out."
       )

    _h(doc, "10.6  DL Explainability Gap", level=2)
    _p(doc,
       "SHAP explainability was applied only to the best tree model (Random Forest, h=1). "
       "Gradient-based attribution for LSTM and GRU was not implemented. This means the "
       "DL results lack the feature-level interpretation provided for the tree models, "
       "making it harder to diagnose why those architectures failed at specific horizons."
       )


# ─────────────────────────────────────────────────────────────────────────────
# Section 11 — Conclusion and Future Work
# ─────────────────────────────────────────────────────────────────────────────

def build_conclusion_section(doc: "Document") -> None:
    """Write Section 11: Conclusion and Future Work."""
    _h(doc, "11. Conclusion and Future Work")

    _h(doc, "11.1  Conclusion", level=2)

    _p(doc,
       "This study set out to determine whether any machine-learning or deep-learning "
       "model can generate statistically significant forecasting improvements over the "
       "random-walk benchmark for Australian ACCU Generic carbon-credit spot prices "
       "at 1-, 7-, and 30-day horizons. The answer is unambiguous: no."
       )

    _p(doc,
       "Six model classes — Random Forest, XGBoost, LightGBM, SARIMAX, LSTM, and GRU — "
       "were evaluated using a rigorous, leakage-free pipeline. Walk-forward "
       "cross-validation governed hyperparameter selection. Forecasting skill was assessed "
       "by the Diebold-Mariano test with Harvey-Leybourne-Newbold small-sample correction. "
       "At every horizon and for every model, the null hypothesis of equal predictive "
       "accuracy versus the random walk could not be rejected at the 5% level. Several "
       "models were significantly worse than the random walk at h = 7 days."
       )

    _p(doc,
       "SHAP analysis of the best tree model (Random Forest, h = 1) identified the "
       "dominant drivers as short-lag momentum (chg_0, the most-recent daily change) "
       "and the HIR cross-market price change (hir_chg), with no staleness feature "
       "(price_moved, days_since_last_move, moves_7d/30d) appearing in the top 8 "
       "by mean absolute SHAP value. These momentum and cross-market signals are real, "
       "but their average impact (≤ 0.023 A$/tonne) is too small relative to the "
       "test-set random-walk RMSE (0.32 A$/tonne) to consistently produce lower "
       "forecast error than simply carrying today's price forward."
       )

    _p(doc,
       "The random walk is therefore the rational benchmark for ACCU spot price "
       "forecasting at daily frequency with the available data and feature set. This is "
       "a credible, defensible result that reflects the near-efficient, event-driven "
       "structure of the Australian carbon market."
       )

    _h(doc, "11.2  Future Work", level=2)

    _p(doc,
       "The limitations identified in Section 10 suggest several directions that could "
       "improve on the current analysis:"
       )

    _bullet(doc,
            "Weekly or move-day aggregation.  Collapsing the data to weekly price "
            "changes, or conditioning the analysis on genuine-move days only, would "
            "substantially increase the effective sample size and reduce the "
            "staleness-induced noise that dominates the current feature space. "
            "This is the highest-priority structural change."
            )
    _bullet(doc,
            "Richer exogenous drivers.  Incorporating EU-ETS prices, Australian "
            "electricity and gas prices, Renewable Energy Certificate supply metrics, "
            "macro variables, and policy-event indicators would directly address the "
            "most likely source of predictable ACCU price variance. A structured "
            "feature-selection step (e.g. LASSO with walk-forward CV) would then "
            "identify which variables carry genuine incremental information."
            )
    _bullet(doc,
            "Event-based modelling.  Given that ACCU price moves appear to be "
            "triggered by discrete policy events (ERF auctions, Safeguard Mechanism "
            "announcements, government procurement decisions), a point-process or "
            "event-study framework may be more appropriate than a pure time-series "
            "regression. A classifier that predicts the probability of a price move "
            "occurring — rather than predicting the continuous change — could be a "
            "more tractable first step."
            )
    _bullet(doc,
            "Probabilistic and interval forecasts.  Point-forecast RMSE comparisons "
            "mask the uncertainty structure of the predictions. Quantile regression "
            "forests, conformalized prediction intervals, or Bayesian approaches would "
            "produce calibrated uncertainty estimates — arguably more useful for "
            "carbon-market participants (who care about tail-risk) than a point forecast."
            )
    _bullet(doc,
            "Longer historical series.  The current dataset begins in January 2018, "
            "giving approximately 6.5 years of data. Extending back to the ERF's "
            "inaugural auctions (2015–2016) and forward to more recent data would "
            "increase the training-set effective sample size and provide more regime "
            "diversity for evaluation."
            )
    _bullet(doc,
            "Nested cross-validation.  A proper multi-window walk-forward evaluation "
            "(multiple non-overlapping test windows, each preceded by a training+tuning "
            "period) would reduce the variance of the skill-score estimate and provide "
            "more reliable evidence about model generalisation across market regimes."
            )
    _bullet(doc,
            "DL attribution.  Applying gradient-based attribution (Integrated Gradients, "
            "GradCAM) or SHAP KernelExplainer to the LSTM and GRU architectures would "
            "close the explainability gap noted in Section 10.6, allowing a direct "
            "comparison of what tree models and recurrent models rely on."
            )

    _p(doc,
       "In summary, the most productive direction is not model complexity but data "
       "enrichment and frequency transformation. The current finding — that no tested "
       "model beats the random walk — is most naturally interpreted as an information "
       "deficit rather than a modelling failure: the features available at daily frequency "
       "do not contain enough signal to outperform the simplest possible baseline."
       )
